import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(transcript_text: str, summary_text: str, evaluation_text: str) -> bytes:
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Transcription Report")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run("Whisper • Summarizer • LLM-as-a-Judge")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() 

    doc.add_heading("Transcription", level=1)
    for line in transcript_text.split("\n"):
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Summary", level=1)
    for line in summary_text.split("\n"):
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("LLM Evaluation", level=1)

    for raw_line in evaluation_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        if ":" in line:
            label, rest = line.split(":", 1)
            p = doc.add_paragraph()
            r_label = p.add_run(label.strip() + ": ")
            r_label.bold = True
            r_body = p.add_run(rest.strip())
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
